"""
Word 文件解析服务
支持 .docx 格式
从 Word 表格和文本中提取教师信息
"""
import re
from docx import Document
from backend.services.field_detector import map_headers, detect_field, normalize_field_name


def parse_word(file_path: str) -> tuple[list[dict], list[str]]:
    """
    解析 Word 文件，提取教师信息

    策略：
    1. 优先从表格中提取（按行或按列）
    2. 从段落文本中用正则提取键值对

    Returns:
        (records, new_fields)
    """
    records = []
    all_new_fields = []

    try:
        doc = Document(file_path)

        # 1. 处理表格
        for table in doc.tables:
            table_records, new_fields = parse_word_table(table)
            records.extend(table_records)
            all_new_fields.extend(new_fields)

        # 2. 处理段落文本（如果没有表格或表格数据不足）
        if not records:
            text_records, new_fields = parse_word_text(doc)
            records.extend(text_records)
            all_new_fields.extend(new_fields)

    except Exception as e:
        raise ValueError(f"Word 解析错误: {str(e)}")

    return records, list(set(all_new_fields))


def parse_word_table(table) -> tuple[list[dict], list[str]]:
    """
    解析 Word 中的表格

    支持两种布局：
    1. 横向表格：第一行是表头，后续行是数据
    2. 纵向表格：左列是字段名，右列是值（一人一表）
    """
    records = []
    new_fields = []

    rows = table.rows
    if len(rows) < 2:
        return records, new_fields

    # 获取所有单元格文本
    cell_texts = []
    for row in rows:
        row_texts = [cell.text.strip() for cell in row.cells]
        cell_texts.append(row_texts)

    # 判断是横向还是纵向表格
    # 横向：第一行包含多个可识别的字段名
    first_row = cell_texts[0]
    detected_count = sum(1 for h in first_row if detect_field(h)[0])

    if detected_count >= 2:
        # 横向表格
        return parse_horizontal_table(cell_texts)
    else:
        # 尝试纵向表格
        return parse_vertical_table(cell_texts)


def parse_horizontal_table(cell_texts: list[list[str]]) -> tuple[list[dict], list[str]]:
    """解析横向表格（第一行为表头）"""
    records = []
    headers = cell_texts[0]
    header_mapping, new_fields = map_headers(headers)

    if not header_mapping:
        return records, new_fields

    for row_texts in cell_texts[1:]:
        record = {}
        has_data = False
        for col_idx, (field_name, is_builtin) in header_mapping.items():
            if col_idx < len(row_texts) and row_texts[col_idx]:
                record[field_name] = row_texts[col_idx]
                has_data = True
        if has_data and record:
            records.append(record)

    return records, new_fields


def parse_vertical_table(cell_texts: list[list[str]]) -> tuple[list[dict], list[str]]:
    """解析纵向表格（左列为字段名，右列为值）"""
    record = {}
    new_fields = []

    for row_texts in cell_texts:
        if len(row_texts) >= 2:
            # 可能有多个键值对在同一行（如 姓名：张三  性别：男）
            i = 0
            while i < len(row_texts) - 1:
                key = row_texts[i]
                value = row_texts[i + 1]
                if key and value:
                    field_name, is_builtin = detect_field(key)
                    if field_name:
                        record[field_name] = value
                        if not is_builtin:
                            new_fields.append(field_name)
                i += 2

    if record:
        return [record], new_fields
    return [], new_fields


def parse_word_text(doc) -> tuple[list[dict], list[str]]:
    """
    从 Word 段落文本中提取教师信息
    支持格式如 "姓名：张三" 或 "姓名: 张三"
    """
    record = {}
    new_fields = []

    # 匹配 "字段名：值" 或 "字段名: 值" 的模式
    pattern = re.compile(r'([^\s:：]+?)\s*[：:]\s*([^\s:：,，;；]+)')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        matches = pattern.findall(text)
        for key, value in matches:
            field_name, is_builtin = detect_field(key)
            if field_name and value:
                record[field_name] = value
                if not is_builtin:
                    new_fields.append(field_name)

    if record:
        return [record], new_fields
    return [], new_fields
